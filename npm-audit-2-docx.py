import logging
import os
import re
from pathlib import Path

from docx import Document

LOGGING_FORMAT = '%(asctime)s | %(levelname)8s | %(message)s'
LOGGING_TIME_AND_DATE_FORMAT = '%d/%b/%Y %a %H:%M:%S %z'

NPM_AUDIT_FILE_EXTENSION = '.txt'
INPUT_DIRECTORY_WITH_REPORTS = './npm'
OUTPUT_FILE_EXTENSION = '.docx'
OUTPUT_DIRECTORY = 'output'


# More checks can be implemented
def is_suitable_npm_audit_file(file_to_check: str) -> bool:
    extension = os.path.splitext(file_to_check)[-1].lower()
    return extension == NPM_AUDIT_FILE_EXTENSION


def get_severity_number(severity: str) -> int:
    if "critical" in severity.lower():
        return 4
    if "high" in severity.lower():
        return 3
    if "medium" in severity.lower() or "moderate" in severity.lower():
        return 2
    if "low" in severity.lower():
        return 1
    else:
        return 0


def process_npm_audit_report_to_docx(npm_audit_file_path: str):
    logging.info(f'Processing npm audit file to docx: {npm_audit_file_path}')

    logging.info('Creating Word document object')
    document = Document()
    document.add_heading(npm_audit_file_path.upper())
    document.add_page_break()

    parsed_vulnerabilities_set = set()

    logging.info(f'Opening npm audit file to parse vulnerabilities: {npm_audit_file_path}')
    with open(npm_audit_file_path) as npm_audit_file:
        logging.info("Reading vulnerabilities information from file")
        for line in npm_audit_file:
            # Split npm audit parseable input on tab
            line_contents = re.split(r'\t+', line.rstrip('\t'))

            # Ensure we do not parse empty lines
            if len(line_contents) >= 6:
                severity_num = get_severity_number(line_contents[2])
                line_tuple = (line_contents[1], line_contents[2], severity_num, line_contents[4],
                              line_contents[5])  # Component, Severity, Severity Num (for sorting only), Type, Advisory URL
                logging.info(f'Adding vulnerabilities: {line_tuple} to set')
                parsed_vulnerabilities_set.add(line_tuple)

    logging.info("Sorting set with vulns in order of severity first")

    sorted_vulnerabilities = sorted(parsed_vulnerabilities_set, key=lambda x: (-x[2], x[0]))  # sort by severity first (desc), then vulnerable component

    logging.info("Preparing docx table with vulnerabilities")

    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Vulnerable Component'
    hdr_cells[1].text = 'Severity'
    hdr_cells[2].text = 'Vulnerability Type'
    hdr_cells[3].text = 'Advisory URL'

    for vuln_tuple in sorted_vulnerabilities:
        row_cells = table.add_row().cells
        row_cells[0].text = vuln_tuple[0]  # Component
        row_cells[1].text = vuln_tuple[1]  # Severity
        row_cells[2].text = vuln_tuple[3]  # Type
        row_cells[3].text = vuln_tuple[4]  # Advisory URL

    output_file = os.path.join(OUTPUT_DIRECTORY, Path(file).with_suffix(OUTPUT_FILE_EXTENSION))
    logging.info(f'Finished processing. Saving to: {output_file}')
    document.save(output_file)


if __name__ == '__main__':
    logging.basicConfig(level=logging.INFO, format=LOGGING_FORMAT, datefmt=LOGGING_TIME_AND_DATE_FORMAT)
    logging.info(f'Creating output directory: {OUTPUT_DIRECTORY}')
    os.makedirs(OUTPUT_DIRECTORY, exist_ok=True)

    # Traverse the input directory recursively
    for root_directory, subdirectories, files_list in os.walk(INPUT_DIRECTORY_WITH_REPORTS):
        for file in files_list:
            file_abs_path = os.path.join(os.path.abspath(root_directory), file)
            if is_suitable_npm_audit_file(file):
                logging.info(f'Found npm audit file: {file_abs_path}')
                process_npm_audit_report_to_docx(file_abs_path)
            else:
                logging.warning(f'Skipping file: {file_abs_path}')
